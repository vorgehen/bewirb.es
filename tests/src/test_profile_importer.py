from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from textx.exceptions import TextXSyntaxError

from src.profile_importer import (
    _load_prompt,
    _strip_code_fences,
    collect_input_files,
    extract_text_from_docx,
    import_profile_documents,
    merge_documents,
    validate_profile_dsl,
)

# Unit-Tests laufen ohne API-Key
pytestmark = pytest.mark.unit


# ─── Text-Extraktion ────────────────────────────────────────────────────────


def _make_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    d = Document()
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(path))


def test_extract_text_from_docx_returns_paragraphs(tmp_path: Path) -> None:
    docx = tmp_path / "cv.docx"
    _make_docx(docx, ["Zeile A", "  Zeile B  ", "", "Zeile C"])
    text = extract_text_from_docx(docx)
    assert "Zeile A" in text
    assert "Zeile B" in text
    assert "Zeile C" in text


def test_extract_text_from_docx_includes_table_content(tmp_path: Path) -> None:
    from docx import Document

    docx = tmp_path / "with_table.docx"
    d = Document()
    d.add_paragraph("Intro")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Zelle1"
    t.rows[0].cells[1].text = "Zelle2"
    d.save(str(docx))
    text = extract_text_from_docx(docx)
    assert "Intro" in text
    assert "Zelle1" in text
    assert "Zelle2" in text


def test_collect_input_files_single_file(tmp_path: Path) -> None:
    docx = tmp_path / "f.docx"
    _make_docx(docx, ["x"])
    files = collect_input_files([docx])
    assert files == [docx]


def test_collect_input_files_recurses_directory(tmp_path: Path) -> None:
    sub = tmp_path / "sub"
    sub.mkdir()
    a = tmp_path / "a.docx"
    b = sub / "b.docx"
    _make_docx(a, ["a"])
    _make_docx(b, ["b"])
    files = collect_input_files([tmp_path])
    assert a in files and b in files


def test_collect_input_files_raises_when_empty(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        collect_input_files([tmp_path])


def test_collect_input_files_raises_for_missing_path(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        collect_input_files([tmp_path / "nicht_vorhanden"])


def test_merge_documents_adds_filename_separators(tmp_path: Path) -> None:
    a = tmp_path / "alpha.docx"
    b = tmp_path / "beta.docx"
    _make_docx(a, ["Inhalt A"])
    _make_docx(b, ["Inhalt B"])
    merged = merge_documents([a, b])
    assert "--- alpha.docx ---" in merged
    assert "--- beta.docx ---" in merged
    assert "Inhalt A" in merged and "Inhalt B" in merged


# ─── Hilfsfunktionen ────────────────────────────────────────────────────────


def test_load_prompt_substitutes_placeholder() -> None:
    prompt = _load_prompt("MEIN-TESTTEXT-123")
    assert "MEIN-TESTTEXT-123" in prompt
    assert "{{profil_text}}" not in prompt


def test_strip_code_fences_removes_block() -> None:
    raw = '```textx\nperson P { title: "X" contact { email: "x@x" } }\n```'
    assert _strip_code_fences(raw).startswith("person P {")


def test_strip_code_fences_passes_through_plain_text() -> None:
    raw = 'person P { title: "X" contact { email: "x@x" } }'
    assert _strip_code_fences(raw) == raw


# ─── DSL-Validierung ────────────────────────────────────────────────────────


_VALID_MINIMAL = """
branche IT { label: "IT" }
person Demo {
    title: "Senior Engineer"
    contact { email: "x@x.de" }
}
"""


def test_validate_profile_dsl_accepts_valid_input() -> None:
    validate_profile_dsl(_VALID_MINIMAL)  # no raise


def test_validate_profile_dsl_rejects_invalid_input() -> None:
    with pytest.raises(TextXSyntaxError):
        validate_profile_dsl("garbage that is no profile DSL")


# ─── Orchestrierung (mit gemocktem Claude) ──────────────────────────────────


def test_import_profile_documents_writes_validated_dsl(tmp_path: Path) -> None:
    docx = tmp_path / "cv.docx"
    _make_docx(docx, ["Robert Risch", "Senior Engineer"])
    output = tmp_path / "result.profile"

    with patch("src.profile_importer.call_claude_for_profile") as mock_call:
        mock_call.return_value = _VALID_MINIMAL
        result = import_profile_documents([docx], output)

    assert result == output
    assert output.exists()
    content = output.read_text(encoding="utf-8")
    assert "person Demo" in content
    mock_call.assert_called_once()
    # Der gemockte Claude-Aufruf bekam den extrahierten Text als Input
    sent_text = mock_call.call_args[0][0]
    assert "Robert Risch" in sent_text


def test_import_profile_documents_raises_when_claude_returns_invalid_dsl(
    tmp_path: Path,
) -> None:
    docx = tmp_path / "cv.docx"
    _make_docx(docx, ["X"])
    output = tmp_path / "out.profile"

    with patch("src.profile_importer.call_claude_for_profile") as mock_call:
        mock_call.return_value = "kein gueltiges DSL"
        with pytest.raises(TextXSyntaxError):
            import_profile_documents([docx], output)

    assert not output.exists()  # bei Fehler keine Datei geschrieben


def test_import_profile_documents_handles_directory_input(tmp_path: Path) -> None:
    sub = tmp_path / "cv_dir"
    sub.mkdir()
    a = sub / "a.docx"
    b = sub / "b.docx"
    _make_docx(a, ["A-Text"])
    _make_docx(b, ["B-Text"])
    output = tmp_path / "merged.profile"

    with patch("src.profile_importer.call_claude_for_profile") as mock_call:
        mock_call.return_value = _VALID_MINIMAL
        import_profile_documents([sub], output)

    # Beide Quellen sind im an Claude gesendeten Text
    sent_text = mock_call.call_args[0][0]
    assert "A-Text" in sent_text and "B-Text" in sent_text


# ─── Optional: echter API-Aufruf (nur mit Marker) ───────────────────────────


@pytest.mark.integration
def test_integration_real_claude_call(tmp_path: Path) -> None:
    """End-to-End mit echter Claude API. Benötigt ANTHROPIC_API_KEY."""
    import os

    if not os.environ.get("ANTHROPIC_API_KEY"):
        pytest.skip("ANTHROPIC_API_KEY nicht gesetzt")

    docx = tmp_path / "mini.docx"
    _make_docx(
        docx,
        [
            "Lebenslauf — Max Beispiel",
            "Senior Software Engineer",
            "Kontakt: max@beispiel.de",
            "Telefon: +49 123 456789",
            "Standort: Köln",
            "",
            "Sprachen: Deutsch (Muttersprache), Englisch (verhandlungssicher)",
            "",
            "Berufserfahrung:",
            "01.2020 – heute, Beispiel GmbH, Bonn, Senior Java Entwickler",
            "  Migration einer Monolith-Anwendung auf Microservices (Java, Spring Boot).",
            "",
            "Ausbildung:",
            "2005-2010, Universität Beispielstadt, Informatik (M.Sc.)",
            "",
            "Zertifikate:",
            "2018 - iSAQB Certified Professional for Software Architecture",
        ],
    )
    output = tmp_path / "imported.profile"
    result = import_profile_documents([docx], output)
    assert result.exists()
    content = result.read_text(encoding="utf-8")
    assert "person" in content
    assert "Max" in content or "Beispiel" in content
