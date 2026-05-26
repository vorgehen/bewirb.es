"""Anonymisierung von Kundendokumenten (Phase 11a, Level 1).

Lokale, regelbasierte Anonymisierung — kein LLM, kein API-Call.
Eingabe-Dokumente bleiben lokal; nur das anonymisierte Extrakt
verlässt diesen Schritt.

Drei Schichten:
1. Standard-Regex: Email, IBAN, IP, URL, Telefon
2. Konfig-Wörterbuch (`anonymize_config.yaml`): gezielte Ersetzungen
3. Public-Domain-Schutz: Begriffe die nie ersetzt werden

Geplant (Level 2+3, später):
- spaCy NER für Organisationen/Personen/Orte
- Optional LLM-Postprocessing
- Strukturiertes Extrakt für `enrich_profile.py --preprocess`
"""

from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path

import yaml
from docx import Document
from pydantic import BaseModel

# ─── Standard-Pattern (Level 1) ────────────────────────────────────────────

_PATTERNS: list[tuple[str, str, str]] = [
    # (kategorie, regex, replacement)
    (
        "email",
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "[EMAIL]",
    ),
    (
        "iban",
        r"\b[A-Z]{2}\d{2}(?:\s?\d{4}){2,7}(?:\s?\d{1,4})?\b",
        "[IBAN]",
    ),
    (
        "ipv4",
        r"\b(?:\d{1,3}\.){3}\d{1,3}\b",
        "[IP]",
    ),
    (
        "url",
        r"\bhttps?://[^\s,;)\]]+",
        "[URL]",
    ),
    (
        "telefon",
        # +49 123 456789, 0123/4567890, (089) 123-456, etc. — defensiv
        r"(?:\+?\d{1,3}[\s\-./]?)?(?:\(?\d{2,5}\)?[\s\-./]?){2,4}\d{2,5}",
        "[TEL]",
    ),
]


# ─── Konfiguration ─────────────────────────────────────────────────────────


class AnonymizeConfig(BaseModel):
    replace: dict[str, str] = {}
    public_domain: list[str] = []
    patterns: list[dict[str, str]] = []

    @classmethod
    def load(cls, path: Path | None) -> AnonymizeConfig:
        if path is None or not path.exists():
            return cls()
        data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        return cls(**data)


# ─── Audit-Log ─────────────────────────────────────────────────────────────


class Replacement(BaseModel):
    original: str
    ersetzung: str
    kategorie: str  # email / iban / ip / url / telefon / replace / pattern
    count: int = 1


class AuditLog(BaseModel):
    replacements: list[Replacement] = []
    public_domain_skipped: list[str] = []

    def add(self, original: str, ersetzung: str, kategorie: str) -> None:
        for r in self.replacements:
            if r.original == original and r.kategorie == kategorie:
                r.count += 1
                return
        self.replacements.append(
            Replacement(original=original, ersetzung=ersetzung, kategorie=kategorie)
        )

    def format(self) -> str:
        lines = ["=== Anonymisierungs-Audit ==="]
        if not self.replacements and not self.public_domain_skipped:
            lines.append("(keine Ersetzungen)")
            return "\n".join(lines)
        by_kat: dict[str, list[Replacement]] = {}
        for r in self.replacements:
            by_kat.setdefault(r.kategorie, []).append(r)
        for kat in sorted(by_kat):
            lines.append(f"\n[{kat}]")
            for r in by_kat[kat]:
                lines.append(f"  {r.original!r:40s} → {r.ersetzung!r}  ({r.count}x)")
        if self.public_domain_skipped:
            lines.append("\n[public-domain (geschützt vor Ersetzung)]")
            for term in self.public_domain_skipped:
                lines.append(f"  {term!r}")
        return "\n".join(lines)


# ─── Pipeline ──────────────────────────────────────────────────────────────


def _is_in_public_domain(text: str, public_domain: list[str]) -> bool:
    """Prüft ob ein potentiell zu ersetzendes Match Teil eines geschützten
    Begriffs ist."""
    return any(pd.lower() in text.lower() for pd in public_domain)


def anonymize_text(text: str, config: AnonymizeConfig | None = None) -> tuple[str, AuditLog]:
    """Wendet alle Anonymisierungs-Regeln auf einen Text an.

    Reihenfolge:
    1. Konfig-Wörterbuch (replace) — längste Keys zuerst, damit
       "Beispiel-Bank Frankfurt" vor "Beispiel-Bank" matcht.
    2. Konfig-Pattern (patterns)
    3. Standard-Pattern (email/iban/ip/url/telefon)

    Public-Domain-Begriffe werden vor jeder Ersetzung geprüft und
    geschützt.
    """
    config = config or AnonymizeConfig()
    audit = AuditLog()

    public_domain_seen: set[str] = set()
    for pd in config.public_domain:
        if re.search(r"\b" + re.escape(pd) + r"\b", text, re.IGNORECASE):
            public_domain_seen.add(pd)
    audit.public_domain_skipped = sorted(public_domain_seen)

    # 1. Konfig-Wörterbuch (replace)
    for key in sorted(config.replace.keys(), key=len, reverse=True):
        value = config.replace[key]
        if _is_in_public_domain(key, config.public_domain):
            continue
        pattern = re.escape(key)
        new_text, n = re.subn(pattern, value, text)
        if n > 0:
            text = new_text
            audit.add(key, value, "replace")

    # 2. Konfig-Pattern
    for entry in config.patterns:
        pat = entry.get("pattern", "")
        repl = entry.get("replacement", "[REDACTED]")
        if not pat:
            continue
        new_text, n = re.subn(pat, repl, text)
        if n > 0:
            text = new_text
            audit.add(pat, repl, "pattern")

    # 3. Standard-Pattern
    for kategorie, pattern, replacement in _PATTERNS:

        def _replace(match: re.Match[str], _r: str = replacement, _k: str = kategorie) -> str:
            original = match.group(0)
            if _is_in_public_domain(original, config.public_domain):
                return original
            audit.add(original, _r, _k)
            return _r

        text = re.sub(pattern, _replace, text)

    return text, audit


# ─── Dokument-Lesen ───────────────────────────────────────────────────────


def extract_text(path: Path) -> str:
    """Liest Text aus einer Datei (.docx oder .txt/.md)."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.rstrip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.rstrip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8")


def collect_files(inputs: Iterable[Path]) -> list[Path]:
    """Sammelt Dateien aus Pfaden (Einzeldatei oder Verzeichnis rekursiv)."""
    files: list[Path] = []
    for inp in inputs:
        if inp.is_dir():
            for ext in (".docx", ".txt", ".md"):
                files.extend(sorted(inp.rglob(f"*{ext}")))
        elif inp.is_file():
            files.append(inp)
        else:
            raise FileNotFoundError(f"Pfad nicht gefunden: {inp}")
    if not files:
        raise FileNotFoundError("Keine lesbaren Dokumente gefunden.")
    return files


def anonymize_documents(
    inputs: list[Path], config_path: Path | None = None
) -> tuple[str, AuditLog]:
    """Hauptfunktion: liest alle Dokumente, konkatieniert, anonymisiert.

    Gibt (anonymisierter_text, audit_log) zurück.
    """
    config = AnonymizeConfig.load(config_path) if config_path else AnonymizeConfig()
    files = collect_files(inputs)
    chunks: list[str] = []
    for f in files:
        chunks.append(f"--- {f.name} ---\n{extract_text(f)}")
    raw = "\n\n".join(chunks)
    return anonymize_text(raw, config)
