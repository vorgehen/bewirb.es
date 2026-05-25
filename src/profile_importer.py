"""Import-Tool: CV/Profil-Dokumente (.docx) → .profile DSL via Claude API.

Datenfluss:
    *.docx → python-docx Text-Extraktion → zusammenführen → Claude API
    → DSL-Text → TextX-Validierung → Datei schreiben
"""

from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path

import anthropic
from anthropic.types import TextBlock
from docx import Document
from textx import metamodel_from_file

PROMPT_FILE = Path(__file__).parent.parent / "prompts" / "import_profile.md"
GRAMMAR_FILE = Path(__file__).parent.parent / "grammar" / "profile.tx"
MODEL = "claude-sonnet-4-6"
MAX_TOKENS = 8000


# ─── Text-Extraktion ────────────────────────────────────────────────────────


def extract_text_from_docx(path: Path) -> str:
    """Liest Absätze und Tabellen-Zellen aus einer .docx-Datei."""
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.rstrip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.rstrip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def collect_input_files(inputs: Iterable[Path]) -> list[Path]:
    """Sammelt .docx-Dateien aus Pfaden — Datei oder Verzeichnis (rekursiv)."""
    files: list[Path] = []
    for inp in inputs:
        if inp.is_dir():
            files.extend(sorted(inp.rglob("*.docx")))
        elif inp.is_file():
            files.append(inp)
        else:
            raise FileNotFoundError(f"Pfad nicht gefunden: {inp}")
    if not files:
        raise FileNotFoundError("Keine .docx-Dateien in den Eingabe-Pfaden gefunden.")
    return files


def merge_documents(files: Iterable[Path]) -> str:
    """Konkateniert den Text mehrerer .docx-Dateien mit Trennlinien."""
    chunks: list[str] = []
    for f in files:
        text = extract_text_from_docx(f)
        if text:
            chunks.append(f"--- {f.name} ---\n{text}")
    return "\n\n".join(chunks)


# ─── Claude-Aufruf ──────────────────────────────────────────────────────────


def _load_prompt(profil_text: str) -> str:
    template = PROMPT_FILE.read_text(encoding="utf-8")
    return template.replace("{{profil_text}}", profil_text)


def _strip_code_fences(text: str) -> str:
    """Entfernt führende/abschließende ```-Code-Fences falls Claude doch welche setzt."""
    text = text.strip()
    fence_re = re.compile(r"^```(?:\w+)?\n(.*?)\n```$", re.DOTALL)
    m = fence_re.match(text)
    if m:
        return m.group(1).strip()
    return text


def _clean_markdown_artifacts(text: str) -> str:
    """Entfernt häufige Markdown-Reste aus Claude-Output.

    - Führende `*` oder `-` direkt vor einem String-Literal (Bullet-Reflex)
    - `**bold**` Markierungen
    """
    # *"..." oder -"..." → "..."
    text = re.sub(r'([:,\[\s])\s*[*-]\s*("[^"]*")', r"\1 \2", text)
    # **text** → text
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text


def call_claude_for_profile(profil_text: str) -> str:
    """Sendet den Text an Claude und gibt den rohen DSL-Output zurück."""
    client = anthropic.Anthropic()
    prompt = _load_prompt(profil_text)
    message = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        messages=[{"role": "user", "content": prompt}],
    )
    block = message.content[0]
    if not isinstance(block, TextBlock):
        raise ValueError(f"Unerwarteter Block-Typ: {type(block)}")
    text = _strip_code_fences(block.text)
    return _clean_markdown_artifacts(text)


# ─── Validierung ────────────────────────────────────────────────────────────


def validate_profile_dsl(dsl_text: str) -> None:
    """Wirft TextXSyntaxError wenn der DSL-Text die Grammatik verletzt."""
    mm = metamodel_from_file(str(GRAMMAR_FILE))
    mm.model_from_str(dsl_text)


# ─── Orchestrierung ─────────────────────────────────────────────────────────


def import_profile_documents(inputs: list[Path], output: Path) -> Path:
    """Konvertiert Eingabe-Dokumente in eine .profile-Datei.

    Schritte: collect → extract → merge → Claude → Draft schreiben → validate → finalize.

    Robustheit: Der rohe Claude-Output wird vor der Validierung als
    `<output>.draft` gespeichert. Bei Validierungsfehler bleibt die Draft-Datei
    bestehen und kann manuell korrigiert werden — der teure API-Aufruf ist
    so nie umsonst.
    """
    files = collect_input_files(inputs)
    merged = merge_documents(files)
    if not merged:
        raise ValueError("Kein extrahierbarer Text in den Eingabe-Dokumenten.")
    dsl = call_claude_for_profile(merged)

    output.parent.mkdir(parents=True, exist_ok=True)
    draft = output.with_suffix(output.suffix + ".draft")
    draft.write_text(dsl, encoding="utf-8")

    validate_profile_dsl(dsl)  # bei Fehler bleibt die Draft-Datei stehen

    output.write_text(dsl, encoding="utf-8")
    draft.unlink(missing_ok=True)
    return output
