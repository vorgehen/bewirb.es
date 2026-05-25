from __future__ import annotations

from pathlib import Path
from typing import Any

import networkx as nx
from docx import Document
from docxtpl import DocxTemplate, Listing

from src.data_loader import Anforderungen, Profil

DEFAULT_TEMPLATE = Path(__file__).parent.parent.parent / "templates" / "profil.docx"


def _clear_and_set(cell: Any, text: str) -> None:
    para = cell.paragraphs[0]
    para.clear()
    para.add_run(text)


def _add_table_loop(
    doc: Document,
    headers: list[str],
    cell_templates: list[str],
) -> None:
    """Fügt eine Tabelle mit Header / Loop-Start / Data / Loop-End hinzu.

    headers gibt die Spalten-Titel. cell_templates muss exakt so viele
    Einträge haben wie headers — jeweils der Jinja-Ausdruck für die Datenzeile.
    Loop-Variable: 'item'.
    """
    cols = len(headers)
    table = doc.add_table(rows=4, cols=cols)
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(table.rows[1].cells[0], "{%tr for item in items %}")
    for i, expr in enumerate(cell_templates):
        _clear_and_set(table.rows[2].cells[i], expr)
    _clear_and_set(table.rows[3].cells[0], "{%tr endfor %}")


def create_default_template(output: Path) -> None:
    """Creates a starter Word template with Jinja2 placeholders for docxtpl.

    Aufbau:
      1. Kopf: Person + Kontakt
      2. Kurzprofil (optional)
      3. Zielrolle (optional)
      4. Schlüsselkompetenzen (zielgruppen-spezifisch geordnet und benannt)
      5. Technologiekompetenz
      6. Projekterfahrung
      7. Werdegang (Festanstellungen)
      8. Ausbildung
      9. Zertifikate
     10. Sprachen
     11. Persönliche Daten (optional)
    """
    doc = Document()

    # ─── 1. Kopf ───────────────────────────────────────────────────────────
    doc.add_heading("{{ person_title }}", level=1)
    doc.add_paragraph("{{ contact_email }}  |  {{ contact_location }}")

    doc.add_paragraph("{%p if contact_phone %}")
    doc.add_paragraph("Tel.: {{ contact_phone }}")
    doc.add_paragraph("{%p endif %}")

    # ─── 2. Kurzprofil ─────────────────────────────────────────────────────
    doc.add_paragraph("{%p if kurzprofil %}")
    doc.add_heading("Kurzprofil", level=2)
    doc.add_paragraph("{{ kurzprofil }}")
    doc.add_paragraph("{%p endif %}")

    # ─── 3. Zielrolle ──────────────────────────────────────────────────────
    doc.add_paragraph("{%p if zielrolle %}")
    doc.add_paragraph("Zielrolle: {{ zielrolle }}")
    doc.add_paragraph("{%p endif %}")

    # ─── 4. Schlüsselkompetenzen (Kategorie + Items) ───────────────────────
    # Tabelle: Kategorie-Label (links) | Aufzählung der Items (rechts)
    doc.add_paragraph("{%p if schluesselkompetenzen_kategorien %}")
    doc.add_heading("Schlüsselkompetenzen", level=2)
    sk_table = doc.add_table(rows=3, cols=2)
    sk_table.style = "Table Grid"
    _clear_and_set(
        sk_table.rows[0].cells[0],
        "{%tr for kat in schluesselkompetenzen_kategorien %}",
    )
    _clear_and_set(sk_table.rows[1].cells[0], "{{ kat.label }}")
    _clear_and_set(sk_table.rows[1].cells[1], "{{ kat.items_str }}")
    _clear_and_set(sk_table.rows[2].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    # ─── 5. Technologiekompetenz ───────────────────────────────────────────
    doc.add_heading("Technologiekompetenz", level=2)
    tech_table = doc.add_table(rows=4, cols=3)
    tech_table.style = "Table Grid"
    for i, label in enumerate(["Technologie", "Kategorie", "Level"]):
        cell = tech_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(tech_table.rows[1].cells[0], "{%tr for tech in technologien %}")
    dr = tech_table.rows[2]
    _clear_and_set(dr.cells[0], "{{ tech.name }}")
    _clear_and_set(dr.cells[1], "{{ tech.category }}")
    _clear_and_set(dr.cells[2], "{{ tech.proficiency }}")
    _clear_and_set(tech_table.rows[3].cells[0], "{%tr endfor %}")

    # ─── 6. Projekterfahrung ───────────────────────────────────────────────
    doc.add_heading("Projekterfahrung", level=2)
    proj_table = doc.add_table(rows=4, cols=3)
    proj_table.style = "Table Grid"
    for i, label in enumerate(["Zeitraum / Auftraggeber / Rolle", "Projekt", "Technologien"]):
        cell = proj_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(proj_table.rows[1].cells[0], "{%tr for p in projekte %}")
    pr = proj_table.rows[2]
    _clear_and_set(pr.cells[0], "{{ p.start }}–{{ p.end }}")
    pr.cells[0].add_paragraph("{{ p.auftraggeber_label }}")
    pr.cells[0].add_paragraph("{{ p.rolle }}")
    _clear_and_set(pr.cells[1], "{{ p.title }}")
    pr.cells[1].add_paragraph("{{ p.description }}")
    pr.cells[1].add_paragraph("{{ p.achievements_str }}")
    _clear_and_set(pr.cells[2], "{{ p.tech_str }}")
    _clear_and_set(proj_table.rows[3].cells[0], "{%tr endfor %}")

    # ─── 7. Werdegang (Festanstellungen) ───────────────────────────────────
    doc.add_paragraph("{%p if werdegang %}")
    doc.add_heading("Beruflicher Werdegang", level=2)
    wd_table = doc.add_table(rows=4, cols=3)
    wd_table.style = "Table Grid"
    for i, label in enumerate(["Zeitraum", "Position", "Arbeitgeber"]):
        cell = wd_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(wd_table.rows[1].cells[0], "{%tr for w in werdegang %}")
    wdr = wd_table.rows[2]
    _clear_and_set(wdr.cells[0], "{{ w.start }}–{{ w.end }}")
    _clear_and_set(wdr.cells[1], "{{ w.titel }}")
    wdr.cells[1].add_paragraph("{{ w.beschreibung }}")
    _clear_and_set(wdr.cells[2], "{{ w.arbeitgeber }}")
    _clear_and_set(wd_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    # ─── 8. Ausbildung ─────────────────────────────────────────────────────
    doc.add_heading("Ausbildung", level=2)
    aus_table = doc.add_table(rows=4, cols=3)
    aus_table.style = "Table Grid"
    for i, label in enumerate(["Zeitraum", "Abschluss", "Institution"]):
        cell = aus_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(aus_table.rows[1].cells[0], "{%tr for a in ausbildungen %}")
    ar = aus_table.rows[2]
    _clear_and_set(ar.cells[0], "{{ a.start }}–{{ a.end }}")
    _clear_and_set(ar.cells[1], "{{ a.title }}")
    ar.cells[1].add_paragraph("{{ a.abschluss }}")
    _clear_and_set(ar.cells[2], "{{ a.institution }}")
    _clear_and_set(aus_table.rows[3].cells[0], "{%tr endfor %}")

    # ─── 9. Zertifikate ────────────────────────────────────────────────────
    doc.add_paragraph("{%p if zertifikate %}")
    doc.add_heading("Zertifikate", level=2)
    zert_table = doc.add_table(rows=4, cols=3)
    zert_table.style = "Table Grid"
    for i, label in enumerate(["Jahr", "Titel", "Aussteller"]):
        cell = zert_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(zert_table.rows[1].cells[0], "{%tr for z in zertifikate %}")
    zr = zert_table.rows[2]
    _clear_and_set(zr.cells[0], "{{ z.jahr }}")
    _clear_and_set(zr.cells[1], "{{ z.titel }}")
    _clear_and_set(zr.cells[2], "{{ z.aussteller }}")
    _clear_and_set(zert_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    # ─── 10. Sprachen ──────────────────────────────────────────────────────
    doc.add_paragraph("{%p if sprachen %}")
    doc.add_heading("Sprachen", level=2)
    spr_table = doc.add_table(rows=4, cols=2)
    spr_table.style = "Table Grid"
    for i, label in enumerate(["Sprache", "Level"]):
        cell = spr_table.rows[0].cells[i]
        _clear_and_set(cell, label)
        cell.paragraphs[0].runs[0].bold = True
    _clear_and_set(spr_table.rows[1].cells[0], "{%tr for s in sprachen %}")
    sr = spr_table.rows[2]
    _clear_and_set(sr.cells[0], "{{ s.bezeichnung }}")
    _clear_and_set(sr.cells[1], "{{ s.level }}")
    _clear_and_set(spr_table.rows[3].cells[0], "{%tr endfor %}")
    doc.add_paragraph("{%p endif %}")

    # ─── 11. Persönliche Daten ─────────────────────────────────────────────
    doc.add_paragraph("{%p if persoenliche_daten %}")
    doc.add_heading("Persönliche Daten", level=2)
    pd_table = doc.add_table(rows=1, cols=2)
    pd_table.style = "Table Grid"
    _clear_and_set(pd_table.rows[0].cells[0], "{{ persoenliche_daten.label }}")
    _clear_and_set(pd_table.rows[0].cells[1], "{{ persoenliche_daten.value }}")
    # Weitere Zeilen mit Loop
    doc.add_paragraph("{%p endif %}")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def _person_node_data(psm: nx.DiGraph[str]) -> dict[str, Any]:
    for _, data in psm.nodes(data=True):
        if data.get("type") == "Person":
            return dict(data)
    return {}


def _build_context(psm: nx.DiGraph[str], profil: Profil, anf: Anforderungen) -> dict[str, Any]:
    # Schlüsselkompetenzen aus dem PSM (zielgruppen-spezifisch geordnet)
    person_node = _person_node_data(psm)
    sk_ordered = person_node.get("schluesselkompetenzen_ordered") or []
    schluesselkompetenzen_kategorien: list[dict[str, Any]] = [
        {
            "key": entry["key"],
            "label": entry["label"],
            "items": entry["items"],
            "items_str": ", ".join(entry["items"]),
        }
        for entry in sk_ordered
    ]

    technologien: list[dict[str, Any]] = []
    for tech in sorted(profil.technologien, key=lambda t: t.years, reverse=True):
        technologien.append(
            {
                "name": tech.name,
                "category": tech.category,
                "proficiency": tech.proficiency,
                "years": str(tech.years),
            }
        )

    projekte: list[dict[str, Any]] = []
    for p in profil.projekte:
        tech_str = ", ".join(t.name for t in p.uses)
        achievements = "\n".join(f"• {a}" for a in p.achievements)
        auftraggeber_label = p.auftraggeber.label or p.auftraggeber.name if p.auftraggeber else ""
        projekte.append(
            {
                "title": p.title,
                "start": p.start,
                "end": p.end,
                "auftraggeber_label": auftraggeber_label,
                "rolle": p.rolle,
                "tech_str": tech_str,
                "description": p.description,
                "achievements_str": Listing(achievements) if achievements else "",
            }
        )

    ausbildungen: list[dict[str, Any]] = [
        {
            "title": a.title,
            "institution": a.institution,
            "start": a.start,
            "end": a.end,
            "abschluss": a.abschluss,
        }
        for a in profil.ausbildungen
    ]

    werdegang: list[dict[str, Any]] = [
        {
            "titel": w.titel,
            "arbeitgeber": w.arbeitgeber,
            "start": w.start,
            "end": w.end,
            "beschreibung": w.beschreibung,
        }
        for w in profil.werdegang
    ]

    zertifikate: list[dict[str, Any]] = [
        {
            "titel": z.titel,
            "aussteller": z.aussteller,
            "jahr": str(z.jahr),
            "url": z.url,
        }
        for z in profil.zertifikate
    ]

    sprachen: list[dict[str, Any]] = [
        {"bezeichnung": s.bezeichnung, "level": s.level} for s in profil.sprachen
    ]

    # Persönliche Daten als einfaches Label/Value-Paar — Template zeigt nur Erstes
    persoenliche_daten = None
    pd = profil.person.persoenlicheDaten
    if pd is not None:
        # Stellen das erste gesetzte Feld dar — für minimales Template-Beispiel.
        # In realen Profilen wird das Template um eine Schleife erweitert.
        for label, value in [
            ("Geburtsdatum", pd.geburtsdatum),
            ("Geburtsort", pd.geburtsort),
            ("Staatsangehörigkeit", pd.staatsangehoerigkeit),
            ("Familienstand", pd.familienstand),
            ("Kinder", pd.kinder),
        ]:
            if value:
                persoenliche_daten = {"label": label, "value": value}
                break

    c = profil.person.contact
    return {
        "person_title": profil.person.title,
        "contact_email": c.email if c else "",
        "contact_phone": c.phone if c else "",
        "contact_location": c.location if c else "",
        "kurzprofil": profil.person.kurzprofil,
        "zielrolle": anf.rolle,
        "schluesselkompetenzen_kategorien": schluesselkompetenzen_kategorien,
        "technologien": technologien,
        "projekte": projekte,
        "werdegang": werdegang,
        "ausbildungen": ausbildungen,
        "zertifikate": zertifikate,
        "sprachen": sprachen,
        "persoenliche_daten": persoenliche_daten,
    }


def generate_word_file(
    psm: nx.DiGraph[str],
    profil: Profil,
    anf: Anforderungen,
    output: Path,
    template: Path = DEFAULT_TEMPLATE,
) -> None:
    doc = DocxTemplate(str(template))
    context = _build_context(psm, profil, anf)
    doc.render(context)
    doc.save(str(output))
