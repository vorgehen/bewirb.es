#!/usr/bin/env python3
"""
Build plan_profile_generator.docx

Ablauf:
  1. Referenz-Dokument mit Header / Footer erstellen
  2. Plan-Markdown vorverarbeiten (```markdown-Fences entfernen)
  3. pandoc: Markdown -> Word (nummerierte Abschnitte, IHV)
  4. Post-Processing: Versionsverzeichnis + Management Summary VOR das IHV einfügen
     (kein eigener Überschriften-Stil -> erscheinen NICHT im IHV)
  5. Aufräumen

Header: "Profile Generator — <Ueberschrift 1>" (STYLEREF, deutsches Word)
Footer: Datum  |  Dateiname  |  Seite n von m
"""

import os
import re
import subprocess  # nosec B404 — required to invoke pandoc
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ── Pfade ─────────────────────────────────────────────────────────────────────
#
# Repo-interne Pfade werden relativ zum Skript-Standort aufgelöst.
# Der Plan-Pfad ist nutzerspezifisch und kann über die Umgebungsvariable
# PROFIL_PLAN_PATH überschrieben werden.
#
REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO_ROOT / "output"

PLAN_SRC = Path(
    os.environ.get(
        "PROFIL_PLAN_PATH",
        Path.home() / ".claude" / "plans" / "ticklish-wishing-moonbeam.md",
    )
)
OUTPUT_DOCX = OUTPUT_DIR / "plan_profile_generator.docx"
REF_DOCX = OUTPUT_DIR / "_reference.docx"
TMP_MD = OUTPUT_DIR / "_plan_tmp.md"

FILENAME_DISPLAY = "plan_profile_generator.docx"
TODAY = datetime.now().strftime("%d.%m.%Y")

# ── Dokument-Metadaten ─────────────────────────────────────────────────────────

#
# WICHTIG: Datum nur für NEU erstellte Versionen verwendet TODAY.
# Bestehende Versionen müssen ihr ursprüngliches Erstellungsdatum
# als String fest behalten, sonst wird es bei jedem Lauf überschrieben.
#
VERSIONS = [
    (
        "1.0",
        "17.05.2026",
        "Profil-Inhaber / Claude Sonnet 4.6",
        "Initiale Version — MDA-Architektur-Showcase & Bewerbungstool, Phasen 0-7",
    ),
    (
        "2.0",
        "24.05.2026",
        "Profil-Inhaber / Claude Sonnet 4.6",
        "Datenbasis-Vision: 3D-Pipeline (PIM × Anforderungen × Zielgruppe), "
        "5 Kompetenz-Kategorien, AI Governance Positioning, AI-Evaluator-Simulation",
    ),
    (
        "2.1",
        "25.05.2026",
        "Profil-Inhaber / Claude Opus 4.7",
        "Tool-Identität als Lifecycle-Tool (Pflegen + Bewerben, ADR-011), "
        "Advisory Layer (Phase 8a, Modi 1+2 in M2), Knowledge Layer Stufe 1 (Phase 7.5), "
        "M2 als Profil-Erstellungs-MVP gestrafft, Karriere-Kapital-Verwaltung als Langzeit-Vision",
    ),
]


def ask_for_new_version() -> str:
    """
    Fragt, ob eine neue Version erstellt werden soll.
    Wenn ja: Eintrag in VERSIONS, Versionsnummer im Dateinamen.
    Gibt den Dateinamen-Stem zurück (ohne .docx).
    """
    print("\nSoll eine neue Version erstellt werden? (j/n): ", end="", flush=True)
    antwort = input().strip().lower()
    if antwort not in ("j", "ja", "y", "yes"):
        # Letzte vorhandene Version bestimmen
        last_version = VERSIONS[-1][0]
        stem = f"plan_profile_generator_v{last_version}"
        print(f"   Keine neue Version. Dateiname: {stem}.docx")
        return stem

    print("Versionsnummer (z.B. 1.1): ", end="", flush=True)
    version = input().strip()
    print("Beschreibung der Änderungen: ", end="", flush=True)
    beschreibung = input().strip()
    print("Autor (Enter = 'Profil-Inhaber / Claude Sonnet 4.6'): ", end="", flush=True)
    autor = input().strip() or "Profil-Inhaber / Claude Sonnet 4.6"

    VERSIONS.append((version, TODAY, autor, beschreibung))
    stem = f"plan_profile_generator_v{version}"
    print(f"   Neue Version {version} eingetragen. Dateiname: {stem}.docx")
    return stem


# Management Summary — Absätze als Liste; **text** wird fett
SUMMARY_PARAGRAPHS = [
    "**Version 2.1 schärft die Tool-Identität (ADR-011):** Das System ist ein "
    "Lifecycle-Tool mit zwei gleichwertigen Aktivitäten — *Pflegen* (Import, Anreichern, "
    "Vervollständigen) und *Bewerben* (Advisory, Positionieren, Generieren). "
    "Beide arbeiten auf der gemeinsamen .profile-Wissensbasis und informieren sich gegenseitig.",
    "**Meilenstein 2 ist das MVP** des Profil-Erstellungs-Systems. "
    "Kritischer Pfad in sieben Phasen: 7.5 Knowledge Layer → 8 Grammatik → 9 Import → "
    "10 Lücken → 8a Advisory (Modi 1+2) → 11 Anreicherung → 11a Anonymisierung. "
    "Done-Kriterium: advise.py --scan angebote/ läuft und das erste "
    "angereicherte Profil ist versendet.",
    "**Knowledge Layer Stufe 1 (Phase 7.5):** Handgefertigte YAML-Wissensbasis als Fundament — "
    "technology_taxonomy, competency_graph, role_profiles, plus opinions.yaml als explizite "
    "Bias-Schicht. Kein ESCO, kein SFIA-Compliance-Anspruch. "
    "Stufen 2-4 (SWEBOK/SFIA-Extraktion, Graph-Modell, dynamische Updates) "
    "sind spätere Meilensteine.",
    "**Advisory Layer (Phase 8a):** Vier Modi — Einzel-Bewerbungsempfehlung (1) und "
    "Portfolio-Scan über mehrere Angebote (2) in M2; Marktpositionierung (3) und "
    "Profil-Diagnose (4) in M3 sobald der Knowledge Layer reif ist. "
    "Schließbarkeits-Bewertung unterscheidet Artikulations-, Zertifizierungs-, Erfahrungs- "
    "und strukturelle Lücken.",
    "**Datenbasis × Zielgruppe → Ausgabe:** Fünf Angebotsstile (Behörde, Consultant, "
    "StartUp, Wissenschaftlich, Standard) plus AIGovernance bestimmen Reihenfolge und "
    "Benennung der fünf Kompetenz-Kategorien "
    "(Methoden · Fach · Technologie · Spezialgebiet · Führung). "
    "Dieselbe Datenbasis liefert für unterschiedliche Zielgruppen "
    "unterschiedlich akzentuierte Profile.",
    "**Meilenstein 3 vertieft die Intelligenz:** AI-Evaluator-Simulation (Claude als "
    "Recruiter-Agent, Score 1-10, Keyword-Analyse, Lücken), Agent-Ready Output via JSON-LD "
    "für AI-Recruiter-Bots, AI Governance als Erstklasse-Kompetenz im Metamodell (Phase 13), "
    "Style Learning (Phase 11b). Voraussetzung: M2 abgeschlossen — echtes Profil existiert, "
    "erste Bewerbung versendet.",
    "**Karriere-Kapital-Verwaltung als Langzeit-Vision** (Ideen-Kapitel): "
    "Das Tool entwickelt sich vom Profil-Erstellungs-System zu einem Werkzeug für "
    "strategisches Karriere-Kapital. Skills, Erfahrungen und Positionierung werden "
    "kontinuierlich dokumentiert und entwickelt — Bewerbungen werden zu einem Anwendungsfall, "
    "nicht zum Zweck.",
    '"Die Invarianten bleiben. Die Werkzeuge werden besser." — '
    "Domäne, Qualität, Architektur und Intention bleiben menschliche Verantwortung. "
    "Der Schwerpunkt verschiebt sich von der Codierung zur Beurteilung von Lösungen. "
    "Erfahrung wird Vorteil, nicht Last.",
]


# ── XML-Hilfsfunktionen ───────────────────────────────────────────────────────


def add_field(para: Any, field_code: str) -> None:
    """Fügt einen Word-Feldcode in einen Absatz ein."""
    r1 = para.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc1)

    r2 = para.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {field_code} "
    r2._r.append(it)

    r3 = para.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)


def style_run(run: Any, size_pt: float = 9.0) -> None:
    run.font.size = Pt(size_pt)


def add_border(para: Any, side: str = "bottom") -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), "AAAAAA")
    pBdr.append(el)
    pPr.append(pBdr)


def set_tabs(para: Any, stops: list[tuple[int, str]]) -> None:
    pPr = para._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    for pos, align in stops:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), align)
        tab.set(qn("w:pos"), str(pos))
        tabs_el.append(tab)
    pPr.append(tabs_el)


def make_page_break_element() -> Any:
    """Gibt ein <w:p>-Element mit Seitenumbruch zurück."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def make_para_element(
    text_parts: str | list[tuple[str, bool]],
    size_pt: float = 11,
    space_after_pt: float = 6,
    bold_title: bool = False,
) -> Any:
    """
    Erstellt ein <w:p>-Element mit gemischten Runs.
    text_parts: Liste von (text, is_bold) oder einfach ein String.
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    # Abstand nach Absatz
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:after"), str(int(space_after_pt * 20)))
    pPr.append(pSpacing)
    p.append(pPr)

    if isinstance(text_parts, str):
        text_parts = [(text_parts, bold_title)]

    for text, is_bold in text_parts:
        if not text:
            continue
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if is_bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(sz)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
    return p


def parse_inline_bold(text: str) -> list[tuple[str, bool]]:
    """Zerlegt einen Text mit **bold** in [(text, is_bold), ...] Tupel."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    result = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            result.append((part[2:-2], True))
        else:
            result.append((part, False))
    return result


# ── Schritt 1: Referenz-Dokument mit Header / Footer ─────────────────────────


def build_reference_doc() -> None:
    print("1) Erstelle Referenz-Dokument (Header/Footer)...")

    raw = subprocess.run(  # nosec B603 B607 — pandoc is a trusted external dependency
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True,
        check=True,
    ).stdout
    REF_DOCX.write_bytes(raw)

    doc = Document(str(REF_DOCX))
    section = doc.sections[0]

    # ── Header ──────────────────────────────────────────────────────────────
    header = section.header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r_prefix = hp.add_run("Profile Generator  —  ")
    style_run(r_prefix)

    # STYLEREF "Überschrift 1" — deutsches Word; \l = vorwärts suchen als Fallback
    add_field(hp, 'STYLEREF "\xdcberschrift 1" \\l \\* MERGEFORMAT')
    for r in hp.runs[1:]:
        style_run(r)

    add_border(hp, "bottom")

    # ── Footer ──────────────────────────────────────────────────────────────
    # Tabstopps: zentriert bei 4513, rechtsbündig bei 9072 Twips
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    set_tabs(fp, [(4513, "center"), (9072, "right")])
    add_border(fp, "top")

    # Links: statisches Datum
    r_date = fp.add_run(TODAY)
    style_run(r_date)

    # Mitte: Dateiname
    r_fn = fp.add_run(f"\t{FILENAME_DISPLAY}")
    style_run(r_fn)

    # Rechts: "Seite n von m"
    r_tab = fp.add_run("\t")
    style_run(r_tab)
    r_s = fp.add_run("Seite ")
    style_run(r_s)
    add_field(fp, "PAGE")
    for r in fp.runs[-3:]:
        style_run(r)
    r_von = fp.add_run(" von ")
    style_run(r_von)
    add_field(fp, "NUMPAGES")
    for r in fp.runs[-3:]:
        style_run(r)

    doc.save(str(REF_DOCX))
    print(f"   -> {REF_DOCX.name}")


# ── Schritt 2: Markdown vorverarbeiten ───────────────────────────────────────


def preprocess() -> None:
    print("2) Vorverarbeitung des Plans (```markdown-Fences entfernen)...")
    content = PLAN_SRC.read_text(encoding="utf-8")
    content = re.sub(
        r"```markdown\n(.*?)```",
        lambda m: m.group(1).rstrip("\n"),
        content,
        flags=re.DOTALL,
    )
    TMP_MD.write_text(content, encoding="utf-8")
    print(f"   -> {TMP_MD.name}")


# ── Schritt 3: pandoc ────────────────────────────────────────────────────────


def run_pandoc() -> None:
    print("3) Starte pandoc...")
    cmd = [
        "pandoc",
        str(TMP_MD),
        "-o",
        str(OUTPUT_DOCX),
        "--reference-doc",
        str(REF_DOCX),
        "--toc",
        "--toc-depth=3",
        "--number-sections",
        "--shift-heading-level-by=-1",
        "--wrap=none",
        "-V",
        "lang=de",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, check=False)  # nosec B603 B607
    if result.returncode != 0:
        print(f"FEHLER:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)
    if result.stderr:
        print(f"   Hinweise: {result.stderr.strip()}")
    print(f"   -> {OUTPUT_DOCX.name} ({OUTPUT_DOCX.stat().st_size // 1024} KB, pandoc)")


# ── Schritt 4: Prologue VOR das IHV einfügen (Post-Processing) ───────────────


def make_table_xml(header_cells: list[str], data_rows: list[list[str]]) -> Any:
    """
    Erstellt ein Word-Tabellen-XML-Element ohne doc.add_table().
    header_cells: Liste von Strings (Kopfzeile)
    data_rows: Liste von Listen von Strings
    Spaltenbreiten automatisch aus Textbreite (9072 Twips = ca. 16 cm).
    """
    # Spaltenbreiten (Twips): Version 900, Datum 1200, Autor 2400, Beschreibung Rest
    col_widths = [900, 1200, 2400, 9072 - 900 - 1200 - 2400]

    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9072")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    all_rows = [header_cells] + data_rows
    for row_idx, row_data in enumerate(all_rows):
        is_header = row_idx == 0
        tr = OxmlElement("w:tr")
        for text, col_w in zip(row_data, col_widths, strict=False):
            tc = OxmlElement("w:tc")
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_w))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            if is_header:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DDDDDD")
                tcPr.append(shd)
            tc.append(tcPr)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            if is_header:
                rPr.append(OxmlElement("w:b"))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")  # 10 pt
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), "20")
            rPr.append(sz)
            rPr.append(szCs)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def add_prologue_before_toc() -> None:
    """
    Öffnet das generierte Dokument und fügt Versionsverzeichnis + Management Summary
    direkt VOR dem Inhaltsverzeichnis ein.
    Kein Überschriften-Stil -> erscheinen NICHT im IHV, werden NICHT nummeriert.
    """
    print("4) Post-Processing: Prologue vor IHV einfügen...")

    doc = Document(str(OUTPUT_DOCX))
    body = doc.element.body

    # TOC finden: pandoc erzeugt ein <w:sdt>-Element mit instrText "TOC ..."
    toc_el = None
    for child in body:
        raw_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if raw_tag == "sdt":
            instr = child.find(".//" + qn("w:instrText"))
            if instr is not None and "TOC" in (instr.text or ""):
                toc_el = child
                break
        if raw_tag == "p":
            instr = child.find(".//" + qn("w:instrText"))
            if instr is not None and "TOC" in (instr.text or ""):
                toc_el = child
                break

    if toc_el is None:
        # Fallback: nach erstem Element (Titel)
        children = list(body)
        toc_el = children[1] if len(children) > 1 else children[0]
        print("   [Hinweis] Kein TOC gefunden — Fallback nach Titel.")
    else:
        print("   TOC gefunden.")

    def before(el: Any) -> None:
        """Fügt el direkt vor dem TOC-Element ein (addprevious = Lesereihenfolge)."""
        toc_el.addprevious(el)

    # addprevious() hängt jedes Element direkt vor toc_el an — Einfüge-Reihenfolge
    # ist damit identisch mit der späteren Lesereihenfolge im Dokument.
    # Gewünschte Lesereihenfolge:
    #   Versionsverzeichnis-Titel → Tabelle → Seitenumbruch →
    #   Management-Summary-Titel → Absätze → Seitenumbruch → TOC

    # ── 1. Versionsverzeichnis ────────────────────────────────────────────────
    before(make_para_element("Versionsverzeichnis", size_pt=16, space_after_pt=12, bold_title=True))
    tbl_xml = make_table_xml(
        ["Version", "Datum", "Autor", "Beschreibung"],
        [[v, d, a, desc] for v, d, a, desc in VERSIONS],
    )
    before(tbl_xml)

    # ── Seitenumbruch ─────────────────────────────────────────────────────────
    before(make_page_break_element())

    # ── 2. Management Summary ─────────────────────────────────────────────────
    before(make_para_element("Management Summary", size_pt=16, space_after_pt=12, bold_title=True))
    for para_text in SUMMARY_PARAGRAPHS:
        parts = parse_inline_bold(para_text)
        before(make_para_element(parts, size_pt=11, space_after_pt=8))

    # ── Seitenumbruch vor TOC ─────────────────────────────────────────────────
    before(make_page_break_element())

    doc.save(str(OUTPUT_DOCX))
    kb = OUTPUT_DOCX.stat().st_size // 1024
    print(f"   -> {OUTPUT_DOCX.name} ({kb} KB, mit Prologue)")


# ── Aufräumen ─────────────────────────────────────────────────────────────────


def cleanup() -> None:
    for f in [TMP_MD, REF_DOCX]:
        if f.exists():
            f.unlink()


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== Build plan_profile_generator.docx ===\n")

    # Versions-Abfrage: aktualisiert VERSIONS und OUTPUT_DOCX / FILENAME_DISPLAY
    stem = ask_for_new_version()
    OUTPUT_DOCX = OUTPUT_DIR / f"{stem}.docx"
    FILENAME_DISPLAY = f"{stem}.docx"
    print()

    build_reference_doc()
    preprocess()
    run_pandoc()
    add_prologue_before_toc()
    cleanup()
    print("\nFertig.")
